from flask import Flask, request, jsonify
import re
from flask_cors import CORS
import io
import os
import json
import docx
import requests

app = Flask(__name__)
# Permitir peticiones desde el frontend local de Vite (por defecto puerto 5173 u otros)
CORS(app, resources={r"/*": {"origins": "*"}})

# ---------------- UTILIDADES ----------------

def clean_text(text):
    if not text:
        return ''
    cleaned_text = re.sub(r'[\n\r"]', ' ', text)
    cleaned_text = cleaned_text.replace(';', ',')
    return cleaned_text.strip()

def extract_year(dp, record=None):
    # 🔥 PRIORIDAD RIS (PY)
    if record and record.get('PY'):
        return record.get('PY')

    # NBIB fallback
    if not dp:
        return ''
    
    match = re.search(r'\d{4}', dp)
    return match.group(0) if match else ''

def extract_month(dp):
    if not dp:
        return ''
    match = re.search(r'(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)', dp)
    return match.group(0) if match else ''

def extract_doi_medline(record):
    doi_fields = ['LID', 'AID']
    for field in doi_fields:
        if field in record:
            entries = record[field] if isinstance(record[field], list) else [record[field]]
            for entry in entries:
                if '[doi]' in entry:
                    return entry.split('[doi]')[0].strip()
    return ''

def construir_url(pmid, doi, record):
    # 1. Prioridad: PMID (PubMed)
    if pmid:
        return f"https://pubmed.ncbi.nlm.nih.gov/{pmid}/"

    # 2. RIS → L2 o UR
    if record:
        return record.get('UR') or record.get('L2') or ''

    return ''

def limpiar_doi(doi):
    if not doi:
        return ''

    doi = doi.lower().strip()

    # eliminar prefijos comunes
    doi = doi.replace('https://doi.org/', '')
    doi = doi.replace('http://doi.org/', '')
    doi = doi.replace('doi:', '')

    # reemplazar caracteres problemáticos
    doi = doi.replace('/', '_')
    doi = doi.replace('.', '')

    return doi

def generar_id_unico(record):
    # 🔥 1. NBIB → PMID
    if record.get('PMID'):
        return record.get('PMID')

    # 🔥 2. RIS → DOI limpio
    doi = limpiar_doi(record.get('DOI', ''))
    if doi:
        return doi

    # fallback (muy raro, pero por seguridad)
    return None

# ---------------- PARSER NBIB ----------------

def parse_nbib(file):
    content = file.read().decode('utf-8')
    records = []
    record = {}
    current_tag = None
    
    for line in content.split('\n'):
        line = line.rstrip('\r\n')
        if not line:
            if record:
                records.append(record)
                record = {}
                current_tag = None
            continue
            
        # Detecta nueva etiqueta (ej: "PMID- ", "TI  - ")
        match = re.match(r'^([A-Z0-9]{2,4})\s*- (.*)', line)
        if match:
            tag = match.group(1).strip()
            value = match.group(2).strip()
            current_tag = tag
            
            if tag in record:
                if isinstance(record[tag], list):
                    record[tag].append(value)
                else:
                    record[tag] = [record[tag], value]
            else:
                record[tag] = value
        elif line.startswith(' ') and current_tag:
            # Línea de continuación (como un Abstract largo)
            val = line.strip()
            if isinstance(record[current_tag], list):
                record[current_tag][-1] += " " + val
            else:
                record[current_tag] += " " + val

    if record:
        records.append(record)

    for rec in records:
        yield {
            "PMID": rec.get('PMID', ''),
            "TI": rec.get('TI', ''),
            "AU": ', '.join(rec.get('AU', []) if isinstance(rec.get('AU'), list) else [rec.get('AU', '')]),
            "DP": rec.get('DP', ''),
            "AB": rec.get('AB', ''),
            "JT": rec.get('JT', ''),
            "PT": ', '.join(rec.get('PT', []) if isinstance(rec.get('PT'), list) else [rec.get('PT', '')]),
            "DOI": extract_doi_medline(rec)
        }

# ---------------- PARSER RIS ----------------

def parse_ris(file):
    content = file.read().decode('utf-8')
    records = content.split("ER  -")

    for raw in records:
        if not raw.strip():
            continue
            
        lines = raw.strip().split("\n")
        record = {}

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Detecta línea con tag RIS (ej: AB  - , TI  -)
            if re.match(r'^[A-Z0-9]{2}  - ', line):
                tag = line[:2]
                value = line[6:].strip()

                # 🔥 SOLO aplicar continuidad a AB
                if tag == "AB":
                    current_tag = "AB"
                else:
                    current_tag = None

                if tag in record:
                    if isinstance(record[tag], list):
                        record[tag].append(value)
                    else:
                        record[tag] = [record[tag], value]
                else:
                    record[tag] = value
            
            else:
                # 🔥 SOLO concatenar si es AB
                if current_tag == "AB":
                    record["AB"] += " " + line

        yield {
            "PMID": record.get('ID', ''),
            "TI": record.get('TI', ''),
            "AU": ', '.join(record.get('AU', []) if isinstance(record.get('AU'), list) else [record.get('AU', '')]),
            "DP": record.get('DA', ''),
            "AB": record.get('AB', ''),
            "JT": record.get('JO', '') or record.get('T2', ''),
            "PT": record.get('TY', ''),
            "DOI": record.get('DO', ''),
            "UR": record.get('L2', '')
        }

# ---------------- NORMALIZADOR ----------------

def construir_salida(records):
    data = []

    for record in records:
        uid = generar_id_unico(record)
        title = clean_text(record.get('TI', ''))
        authors = clean_text(record.get('AU', ''))
        dp = record.get('DP', '')

        year = extract_year(dp)
        month = extract_month(dp)

        pmid = record.get('PMID', '')
        doi = record.get('DOI', '')

        url = construir_url(pmid, doi, record)

        data.append({
            "PMID": uid,
            "TI": title,
            "AU": authors,
            "YYYY": year,
            "MMM": month,
            "AB": clean_text(record.get('AB', '')),
            "JT": clean_text(record.get('JT', '')),
            "PT": clean_text(record.get('PT', '')),
            "DP": dp,
            "DOI": doi,
            "URL": url
        })

    return data

# ---------------- ROUTE ----------------

@app.route('/convert', methods=['POST'])
def convert():
    if 'file' not in request.files:
        return jsonify({"error": "No se ha subido un archivo"}), 400

    file = request.files['file']

    if file.filename == '':
        return jsonify({"error": "Archivo no válido"}), 400

    try:
        filename = file.filename.lower()

        if filename.endswith('.nbib'):
            records = parse_nbib(file)
        elif filename.endswith('.ris'):
            records = parse_ris(file)
        else:
            return jsonify({"error": "Formato no soportado"}), 400

        data = construir_salida(records)
        return jsonify({"csvData": data})

    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/save-atc', methods=['POST'])
def save_atc():
    try:
        data = request.json
        if not isinstance(data, list):
            return jsonify({"error": "El cuerpo de la petición debe ser una lista de registros ATC"}), 400
        
        # Guardar en src/data/atc_data.json
        base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        target_path = os.path.join(base_dir, 'src', 'data', 'atc_data.json')
        
        # Asegurarnos de que el directorio existe
        os.makedirs(os.path.dirname(target_path), exist_ok=True)
        
        with open(target_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
            
        return jsonify({"success": True, "count": len(data)})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/save-eurd', methods=['POST'])
def save_eurd():
    try:
        data = request.json
        if not isinstance(data, list):
            return jsonify({"error": "El cuerpo de la petición debe ser una lista de registros EURD"}), 400
        
        # Guardar en src/data/eurd_data.json
        base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        target_path = os.path.join(base_dir, 'src', 'data', 'eurd_data.json')
        
        os.makedirs(os.path.dirname(target_path), exist_ok=True)
        
        with open(target_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
            
        return jsonify({"success": True, "count": len(data)})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/save-productos/<filename>', methods=['POST'])
def save_productos(filename):
    try:
        data = request.json
        if not isinstance(data, list):
            return jsonify({"error": "El cuerpo de la petición debe ser una lista de productos"}), 400
        
        # Limpiar filename por seguridad
        import werkzeug.utils
        secure_name = werkzeug.utils.secure_filename(filename)
        
        base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        target_path = os.path.join(base_dir, 'src', 'data', secure_name)
        
        os.makedirs(os.path.dirname(target_path), exist_ok=True)
        
        with open(target_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
            
        return jsonify({"success": True, "count": len(data)})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/generate-docx', methods=['POST'])
def generate_docx():
    try:
        from flask import send_file
        import tempfile
        
        body = request.get_json()
        if not body:
            return jsonify({"error": "No se proporcionaron datos."}), 400
            
        replacements = body.get('replacements', {})
        tables_data = body.get('tables', {})
        template_filename = body.get('templatePath')
        
        if not template_filename:
            template_filename = 'FV-POE-07-F02 Formato de IPS PF.docx'
        
        base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        template_dir = os.path.join(base_dir, 'src', 'config', 'clienteTemplate')
        template_path = os.path.join(template_dir, template_filename)
        
        if not os.path.exists(template_path):
            # Attempt to find the file if it has a prefix (e.g. from multer)
            found = False
            if os.path.exists(template_dir):
                for f in os.listdir(template_dir):
                    if f.endswith(template_filename):
                        template_path = os.path.join(template_dir, f)
                        found = True
                        break
            if not found:
                return jsonify({"error": f"No se encontró la plantilla en {template_path}"}), 404
            
        doc = docx.Document(template_path)
        
        def replace_text_in_runs(runs, key, val):
            for run in runs:
                if key in run.text:
                    run.text = run.text.replace(key, str(val))
                    return True
            
            for i in range(len(runs)):
                for j in range(i+1, len(runs)+1):
                    combined = "".join(r.text for r in runs[i:j])
                    if key in combined:
                        new_combined = combined.replace(key, str(val))
                        runs[i].text = new_combined
                        for k in range(i+1, j):
                            runs[k].text = ""
                        return True
            return False

        def replace_text_in_paragraph(p):
            for key, val in replacements.items():
                if key in p.text:
                    while key in p.text:
                        replaced = replace_text_in_runs(p.runs, key, val)
                        if not replaced:
                            break

        # --- TABLAS DINAMICAS ---
        tables_data = body.get('tables', {})
        
        import copy
        
        for table in doc.tables:
            for row_idx, row in enumerate(table.rows):
                if not row.cells:
                    continue
                
                # Check if this row is a template row containing a table placeholder
                cell_text = row.cells[0].text.strip()
                placeholder = None
                for p in tables_data.keys():
                    if p in cell_text:
                        placeholder = p
                        break
                        
                if placeholder:
                    rows_data = tables_data[placeholder]
                    template_row = row
                    
                    insert_idx = row_idx + 1
                    current_tr = template_row._tr
                    
                    for item in rows_data:
                        # Clonar la fila XML
                        new_tr = copy.deepcopy(template_row._tr)
                        current_tr.addnext(new_tr)
                        current_tr = new_tr
                        
                        new_row = table.rows[insert_idx]
                        
                        # Rellenar celdas basadas en los values del diccionario manteniendo estilo de fuente
                        values = list(item.values())
                        for i, val in enumerate(values):
                            if i < len(new_row.cells):
                                cell = new_row.cells[i]
                                p = cell.paragraphs[0] if len(cell.paragraphs) > 0 else cell.add_paragraph()
                                if p.runs:
                                    p.runs[0].text = str(val)
                                    for r in p.runs[1:]:
                                        r.text = ""
                                else:
                                    p.text = str(val)
                                    
                        insert_idx += 1
                    
                    # Eliminar la fila plantilla original
                    template_row._element.getparent().remove(template_row._element)
                    break # Pasamos a la siguiente tabla
                    
        # --- FIN TABLAS DINAMICAS ---

        # --- SECCIONES DINÁMICAS ---
        sections_data = body.get('sections', {})
        
        def render_benefit_characterization(paragraph, data):
            indicaciones = data.get('indicaciones', [])
            for ind in indicaciones:
                paragraph.insert_paragraph_before('', style='Normal')
                paragraph.insert_paragraph_before(ind.get('title', ''), style='Heading 3')
                paragraph.insert_paragraph_before('', style='Normal')
                for article in ind.get('articles', []):
                    p_title = paragraph.insert_paragraph_before('', style='Normal')
                    p_title.add_run(article.get('title', '')).bold = True
                    summary = article.get('summary', '')
                    if summary:
                        paragraph.insert_paragraph_before(summary, style='Normal')

        def render_benefit_risk_context(paragraph, data):
            indicaciones = data.get('indicaciones', [])
            for ind in indicaciones:
                paragraph.insert_paragraph_before('', style='Normal')
                paragraph.insert_paragraph_before(ind.get('title', ''), style='Heading 3')
                paragraph.insert_paragraph_before('', style='Normal')

        def render_benefit_risk_analysis(paragraph, data):
            indicaciones = data.get('indicaciones', [])
            ifa = data.get('ifa', 'IFA')
            brand = data.get('brand', 'Marca')
            for ind in indicaciones:
                title = ind.get('title', '')
                paragraph.insert_paragraph_before('', style='Normal')
                paragraph.insert_paragraph_before(title, style='Heading 3')
                paragraph.insert_paragraph_before('', style='Normal')
                
                dose = ind.get('recommendedDose', 'Dosis')
                text = f"{brand} ({ifa}) ha demostrado utilidad en el tratamiento de {title} a una dosis recomendada de {dose}."
                paragraph.insert_paragraph_before(text, style='Normal')
                
                global_analysis = ind.get('globalAnalysis', '')
                if global_analysis:
                    paragraph.insert_paragraph_before(global_analysis, style='Normal')

        paragraphs_to_remove = []
        for p in doc.paragraphs:
            for sec_placeholder, sec_info in sections_data.items():
                if sec_placeholder in p.text:
                    renderer = sec_info.get('renderer')
                    data = sec_info.get('data', {})
                    
                    try:
                        if renderer == 'benefit_characterization':
                            render_benefit_characterization(p, data)
                        elif renderer == 'benefit_risk_context':
                            render_benefit_risk_context(p, data)
                        elif renderer == 'benefit_risk_analysis':
                            render_benefit_risk_analysis(p, data)
                    except Exception as e:
                        print(f"Error rendering section {sec_placeholder}: {e}")
                        
                    paragraphs_to_remove.append(p)
                    
        for p in paragraphs_to_remove:
            p._element.getparent().remove(p._element)
            
        # --- FIN SECCIONES DINÁMICAS ---

        # Reemplazar en párrafos principales del cuerpo
        for p in doc.paragraphs:
            replace_text_in_paragraph(p)
            
        # Reemplazar en tablas del cuerpo
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        replace_text_in_paragraph(p)
                        
        # Reemplazar en encabezados y pies de página de todas las secciones
        for section in doc.sections:
            # Cabeceras (Principal, Primera Página, Páginas Pares)
            for header in [section.header, section.first_page_header, section.even_page_header]:
                if header:
                    for p in header.paragraphs:
                        replace_text_in_paragraph(p)
                    for table in header.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for p in cell.paragraphs:
                                    replace_text_in_paragraph(p)
            # Pies de página
            for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                if footer:
                    for p in footer.paragraphs:
                        replace_text_in_paragraph(p)
                    for table in footer.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for p in cell.paragraphs:
                                    replace_text_in_paragraph(p)
                                    
        # Guardar en un archivo temporal en disco de forma segura
        temp_dir = tempfile.gettempdir()
        temp_out = os.path.join(temp_dir, f"IPS_Generado_{os.getpid()}.docx")
        doc.save(temp_out)
        
        # Enviar archivo de descarga al frontend
        return send_file(
            temp_out,
            as_attachment=True,
            download_name="IPS_Generado.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

# ---------------- RUTA: PROCESAR FICHA TÉCNICA (FT FIXED) ----------------

# Carpeta base para salidas relativas al directorio del backend
BASE_OUTPUT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "word_to_md_output")

def sanitize_folder_name(name):
    """Elimina caracteres no permitidos en nombres de carpeta de Windows."""
    import re
    name = name.strip()
    name = re.sub(r'[\\/:*?"<>|]', '_', name)
    return name or "sin_nombre"

def convert_doc_to_docx(doc_path, output_path):
    """Convierte un .doc a .docx usando mammoth (extrae contenido a HTML y luego a docx con python-docx)."""
    import mammoth
    import docx as python_docx
    from docx.shared import Pt

    with open(doc_path, "rb") as f:
        result = mammoth.convert_to_html(f)
    html = result.value

    # Crear un docx básico con el HTML convertido (texto plano, sin formato rico)
    document = python_docx.Document()
    from bs4 import BeautifulSoup
    soup = BeautifulSoup(html, "html.parser")
    for elem in soup.find_all(True):
        if elem.name in ["h1", "h2", "h3", "h4", "h5", "h6"]:
            level = int(elem.name[1])
            p = document.add_heading(elem.get_text(), level=level)
        elif elem.name == "p":
            document.add_paragraph(elem.get_text())
        elif elem.name == "li":
            document.add_paragraph(elem.get_text(), style="List Bullet")
    document.save(output_path)

def apply_headers_to_bold(md_content):
    """
    Transforma el contenido Markdown:
    - Convierte encabezados # (1-6 niveles) en **negritas**
    - Elimina también '#' sueltos sin texto
    - Elimina el marcador '>' de blockquotes al inicio de línea (evita tabulación)
    - Respeta bloques de código (``` ... ```) sin modificarlos
    - Maneja correctamente saltos de línea Windows (\r\n)
    """
    import re
    lines = md_content.split('\n')
    result = []
    in_code_block = False

    for line in lines:
        # Normalizar fin de línea Windows
        line = line.rstrip('\r')

        # Detectar inicio/fin de bloque de código
        if re.match(r'^\s*```', line):
            in_code_block = not in_code_block
            result.append(line)
            continue

        # Dentro de bloque de código: no tocar nada
        if in_code_block:
            result.append(line)
            continue

        # Eliminar '>' de blockquotes al inicio de línea (uno o más niveles)
        line = re.sub(r'^(>\s*)+', '', line)

        # Convertir encabezados # a **negritas**
        # Captura: 1-6 '#', espacio opcional, texto opcional
        match = re.match(r'^(#{1,6})\s*(.*)', line)
        if match:
            text = match.group(2).strip()
            if text:
                # Quitar cualquier formato Markdown inline residual del encabezado
                # (por ej. **texto** que pandoc puede dejar en encabezados)
                result.append(f'**{text}**')
            else:
                result.append('')  # Encabezado vacío → línea en blanco
        else:
            result.append(line)

    return '\n'.join(result)

@app.route('/api/process-ft', methods=['POST'])
def process_ft():
    """
    Ruta para procesar fichas técnicas Word.
    Recibe: archivos Word (.doc/.docx) + clientName + folderNames (uno por archivo).
    Flujo por archivo:
      1. Si .doc → convierte a .docx con mammoth
      2. Pandoc: .docx → .md (extrae imágenes a media/)
      3. Regex: # encabezados → **negritas**
      4. Pandoc: .md → .docx (reinyecta imágenes)
      5. Guarda todo en la carpeta destino (md + media + docx final)
    """
    from flask import send_file
    import subprocess
    import tempfile
    import shutil

    client_name = request.form.get('clientName', 'Cliente').strip()
    client_folder = sanitize_folder_name(client_name)

    files = request.files.getlist('files[]')
    folder_names = request.form.getlist('folderNames[]')

    if not files:
        return jsonify({"error": "No se recibieron archivos"}), 400

    results = []

    for idx, file in enumerate(files):
        original_filename = file.filename
        folder_name_raw = folder_names[idx] if idx < len(folder_names) else os.path.splitext(original_filename)[0]
        folder_name = sanitize_folder_name(folder_name_raw)

        # Carpeta destino final
        dest_dir = os.path.join(BASE_OUTPUT_DIR, client_folder, folder_name)
        os.makedirs(dest_dir, exist_ok=True)

        # Usar directorio temporal del sistema para trabajar
        tmp_dir = tempfile.mkdtemp(prefix="ft_fixed_")

        try:
            # ── Paso 1: Guardar archivo subido al directorio temporal ──────────────
            ext = os.path.splitext(original_filename)[1].lower()
            tmp_uploaded = os.path.join(tmp_dir, original_filename)
            file.save(tmp_uploaded)

            # ── Paso 2: Convertir .doc → .docx si es necesario ───────────────────
            if ext == '.doc':
                tmp_docx = os.path.join(tmp_dir, os.path.splitext(original_filename)[0] + '.docx')
                convert_doc_to_docx(tmp_uploaded, tmp_docx)
            else:
                tmp_docx = tmp_uploaded

            # ── Paso 3: Pandoc docx → md (extrae imágenes) ──────────────────────
            # IMPORTANTE: usar rutas relativas para --extract-media para evitar
            # problemas con espacios/caracteres especiales en rutas absolutas.
            # Pandoc crea: tmp_dir/media/imageN.png
            # El .md generado referencia: media/imageN.png (ruta relativa)
            base_name = os.path.splitext(original_filename)[0]
            md_filename = f"{base_name}.md"  # Mismo nombre que el original
            tmp_md = os.path.join(tmp_dir, md_filename)
            docx_basename = os.path.basename(tmp_docx)

            pandoc_to_md = subprocess.run(
                [
                    'pandoc', docx_basename,          # Ruta relativa al cwd
                    '-o', md_filename,                # Ruta relativa al cwd
                    '--extract-media=.',              # Extrae imágenes en ./media/ (relativo al cwd)
                    '--wrap=none',                    # Sin saltos de línea artificiales
                    '--to=markdown+raw_html',         # Tablas complejas se emiten como HTML nativo
                    '--markdown-headings=atx',        # Encabezados con # (estilo ATX)
                ],
                capture_output=True, text=True, cwd=tmp_dir
            )

            if pandoc_to_md.returncode != 0:
                results.append({
                    "filename": original_filename,
                    "status": "error",
                    "error": f"Pandoc (docx→md) falló: {pandoc_to_md.stderr}"
                })
                continue

            # ── Paso 4: Aplicar regex # → **negritas** ───────────────────────────
            with open(tmp_md, 'r', encoding='utf-8') as f_md:
                md_content = f_md.read()

            md_transformed = apply_headers_to_bold(md_content)

            with open(tmp_md, 'w', encoding='utf-8') as f_md:
                f_md.write(md_transformed)

            # ── Paso 5: MD → HTML → DOCX (tablas HTML → tablas Word nativas) ────────
            #
            # Por qué dos pasos:
            #   - pandoc md→docx con --from=markdown+raw_html trata <table> como
            #     texto opaco y lo vuelca como texto plano en el DOCX.
            #   - Cuando pandoc lee un .html (--from=html), SÍ interpreta los
            #     <table> y los convierte en tablas nativas de Word.
            #   Esto replica exactamente el proceso manual: Markdown Preview Enhanced
            #   renderizaba el MD a HTML visual → copiar al portapapeles → pegar en Word.
            #
            output_docx_name = f"{base_name} FT FIXED.docx"
            html_filename = f"{base_name}.html"
            tmp_output_docx = os.path.join(tmp_dir, output_docx_name)

            # 5a: MD → HTML (pandoc renderiza las tablas HTML como HTML válido)
            pandoc_md_to_html = subprocess.run(
                [
                    'pandoc', md_filename,
                    '-o', html_filename,
                    '--from=markdown+raw_html',   # Lee Markdown + bloques HTML crudos
                    '--to=html',
                    '--standalone',               # Genera HTML completo con <head>/<body>
                    '--embed-resources',          # Incrusta imágenes como base64 en el HTML
                ],
                capture_output=True, text=True, cwd=tmp_dir
            )

            if pandoc_md_to_html.returncode != 0:
                results.append({
                    "filename": original_filename,
                    "status": "error",
                    "error": f"Pandoc (md→html) falló: {pandoc_md_to_html.stderr}"
                })
                continue

            # 5b: HTML → DOCX (pandoc convierte <table> en tablas nativas de Word)
            pandoc_html_to_docx = subprocess.run(
                [
                    'pandoc', html_filename,
                    '-o', output_docx_name,
                    '--from=html',                # Lee HTML completo (con tablas reales)
                ],
                capture_output=True, text=True, cwd=tmp_dir
            )

            if pandoc_html_to_docx.returncode != 0:
                results.append({
                    "filename": original_filename,
                    "status": "error",
                    "error": f"Pandoc (html→docx) falló: {pandoc_html_to_docx.stderr}"
                })
                continue


            # ── Paso 6: Copiar archivos al destino final ──────────────────────────
            # 6a. Documento original subido por el usuario
            dest_original = os.path.join(dest_dir, original_filename)
            shutil.copy2(tmp_uploaded, dest_original)

            # 6b. Word procesado (FT FIXED)
            dest_docx = os.path.join(dest_dir, output_docx_name)
            shutil.copy2(tmp_output_docx, dest_docx)

            # 6c. Markdown transformado (mismo nombre que el original)
            dest_md = os.path.join(dest_dir, md_filename)
            shutil.copy2(tmp_md, dest_md)

            # 6d. Carpeta de imágenes (si existe, en tmp_dir/media/)
            media_dir_tmp = os.path.join(tmp_dir, 'media')
            if os.path.isdir(media_dir_tmp):
                dest_media = os.path.join(dest_dir, 'media')
                if os.path.exists(dest_media):
                    shutil.rmtree(dest_media)
                shutil.copytree(media_dir_tmp, dest_media)

            results.append({
                "filename": original_filename,
                "outputFilename": output_docx_name,
                "outputPath": dest_docx,
                "mdPath": dest_md,
                "destDir": dest_dir,
                "status": "success"
            })

        except Exception as e:
            import traceback
            traceback.print_exc()
            results.append({
                "filename": original_filename,
                "status": "error",
                "error": str(e)
            })
        finally:
            # Limpiar el directorio temporal del sistema
            try:
                shutil.rmtree(tmp_dir, ignore_errors=True)
            except Exception:
                pass

    return jsonify({"results": results})


@app.route('/api/download-ft', methods=['GET'])
def download_ft():
    """Sirve un archivo procesado para descarga desde la ruta de destino."""
    from flask import send_file
    file_path = request.args.get('path', '')

    if not file_path:
        return jsonify({"error": "Ruta no especificada"}), 400

    # Seguridad: solo permitir archivos dentro de la carpeta base
    abs_path = os.path.abspath(file_path)
    abs_base = os.path.abspath(BASE_OUTPUT_DIR)

    if not abs_path.startswith(abs_base):
        return jsonify({"error": "Acceso denegado"}), 403

    if not os.path.isfile(abs_path):
        return jsonify({"error": "Archivo no encontrado"}), 404

    return send_file(
        abs_path,
        as_attachment=True,
        download_name=os.path.basename(abs_path),
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )



@app.route('/api/save-listado-ips', methods=['POST'])
def save_listado_ips():
    try:
        data = request.json
        if not isinstance(data, list):
            return jsonify({"error": "El cuerpo de la petición debe ser una lista de registros IPS"}), 400
        
        # Guardar en src/data/listado_ips.json
        base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        target_path = os.path.join(base_dir, 'src', 'data', 'listado_ips.json')
        
        os.makedirs(os.path.dirname(target_path), exist_ok=True)
        
        with open(target_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
            
        return jsonify({"success": True, "count": len(data)})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# ---------------- CLAUDE API ----------------

@app.route('/api/claude', methods=['POST'])
def call_claude():
    try:
        data = request.json
        api_key = data.get('api_key')
        prompt = data.get('prompt')
        
        if not api_key or not prompt:
            return jsonify({"error": "api_key y prompt son requeridos"}), 400
            
        headers = {
            "x-api-key": api_key,
            "anthropic-version": "2023-06-01",
            "content-type": "application/json"
        }
        
        model = data.get('model', 'claude-3-haiku-20240307')
        
        payload = {
            "model": model,
            "max_tokens": 2048,
            "messages": [
                {"role": "user", "content": prompt}
            ]
        }
        
        response = requests.post(
            "https://api.anthropic.com/v1/messages",
            headers=headers,
            json=payload
        )
        
        if response.status_code != 200:
            return jsonify({"error": f"Error de Anthropic: {response.text}"}), response.status_code
            
        res_data = response.json()
        text_response = res_data.get('content', [{}])[0].get('text', '')
        
        return jsonify({"response": text_response})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# ---------------- OPENAI API ----------------

@app.route('/api/openai', methods=['POST'])
def call_openai():
    try:
        data = request.json
        api_key = data.get('api_key')
        prompt = data.get('prompt')
        
        if not api_key or not prompt:
            return jsonify({"error": "api_key y prompt son requeridos"}), 400
            
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }
        
        model = data.get('model', 'gpt-4o-mini')
        
        payload = {
            "model": model,
            "temperature": 0.2,
            "messages": [
                {"role": "user", "content": prompt}
            ]
        }
        
        response = requests.post(
            "https://api.openai.com/v1/chat/completions",
            headers=headers,
            json=payload
        )
        
        if response.status_code != 200:
            return jsonify({"error": f"Error de OpenAI: {response.text}"}), response.status_code
            
        res_data = response.json()
        text_response = res_data.get('choices', [{}])[0].get('message', {}).get('content', '')
        
        return jsonify({"response": text_response})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# ---------------- MAIN ----------------

if __name__ == '__main__':
    app.run(debug=True, host="127.0.0.1", port=5000)
